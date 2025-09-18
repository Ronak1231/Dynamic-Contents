import streamlit as st
import google.generativeai as genai
from PIL import Image
from docx import Document
import io

# --- Page Configuration ---
st.set_page_config(
    page_title="Dynamic Marketing Content Generator",
    page_icon="🚀",
    layout="wide"
)

# --- API Configuration ---
# Use Streamlit's secrets management to get the API key
try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
except Exception as e:
    st.error(f"Error configuring Gemini API: {e}. Make sure your API key is set in Streamlit secrets.", icon="🔑")
    st.stop()

# --- Helper Function to create a downloadable Word document ---
def create_word_document(content):
    """Creates a Word document from the generated text and returns it as a bytes object."""
    doc = Document()
    doc.add_heading('Dynamic Marketing Content', level=1)
    
    # Simple parsing based on markdown headers
    lines = content.split('\n')
    for line in lines:
        if line.startswith('## '):
            # Add headings without the '## '
            doc.add_heading(line[3:].strip(), level=2)
        elif line.strip(): # Avoid adding empty paragraphs
            doc.add_paragraph(line)
            
    # Save document to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Main App Interface ---
st.title("🚀 Dynamic Marketing Content Generator")
st.markdown("Upload a product image and name to instantly generate compelling marketing content for all your platforms.")

# --- Session State Initialization ---
# This helps 'remember' the generated content across reruns (e.g., after clicking a button)
if "generated_content" not in st.session_state:
    st.session_state.generated_content = None

# --- Input Fields in Two Columns ---
col1, col2 = st.columns(2)

with col1:
    product_name = st.text_input("Enter Product Name:", placeholder="e.g., 'Aura Smart Watch'")
    product_desc = st.text_area("Optional: Add a brief description:", placeholder="e.g., 'A sleek, waterproof smartwatch with 14-day battery life and advanced health tracking.'")
    
with col2:
    uploaded_image = st.file_uploader("Upload Product Image", type=["jpg", "jpeg", "png"])
    if uploaded_image:
        image = Image.open(uploaded_image)
        st.image(image, caption="Your Uploaded Product", use_column_width=True)

# --- Generate Button and Processing ---
if st.button("✨ Generate Marketing Content", type="primary", use_container_width=True):
    # --- Input Validation ---
    if not uploaded_image or not product_name:
        st.warning("Please upload an image and enter a product name.", icon="⚠️")
    else:
        with st.spinner("🤖 Gemini is crafting your content... Please wait."):
            try:
                # --- Prepare the Prompt for Gemini ---
                # We ask for Markdown output for easier, more reliable parsing.
                prompt = f"""
                You are a world-class marketing assistant. Based on the product image and details provided, generate the following content in clear, separate sections using Markdown formatting (using '##' for headings):

                Product Name: {product_name}
                Product Description: {product_desc}

                Generate the following assets:
                
                ## 📢 Ad Copy
                (Create 3 short, catchy ad copy lines, max 15 words each)

                ## 💼 LinkedIn Post
                (A professional post of 100-150 words, highlighting features and benefits for a business audience)

                ## 📧 Marketing Email
                (A formal marketing email of 200-250 words, including a compelling subject line)

                ## 📱 Social Media Captions
                (Create separate, engaging captions for Instagram, Twitter/X, and Facebook. Include relevant hashtags for each.)
                """

                model = genai.GenerativeModel("gemini-1.5-flash-latest")
                image_for_model = Image.open(uploaded_image)
                
                response = model.generate_content([prompt, image_for_model])
                
                # Store the successful response in session state
                st.session_state.generated_content = response.text
                st.success("Content generated successfully!", icon="✅")

            except Exception as e:
                st.error(f"An error occurred: {e}", icon="❌")


# --- Display Output ---
if st.session_state.generated_content:
    st.markdown("---")
    st.subheader("Your Generated Content")

    # Parse the content based on Markdown headers
    content = st.session_state.generated_content
    sections = content.split('## ')
    
    # Create tabs for each section for a cleaner UI
    tab_titles = [s.split('\n')[0].strip() for s in sections if s]
    
    if tab_titles:
        tabs = st.tabs(tab_titles)
        for i, section_content in enumerate(sections[1:]): # Skip the empty first element
            with tabs[i]:
                st.markdown(section_content)
                # Add a copy button for each section
                st.code(section_content.strip(), language=None)
                st.button(f"Copy {tab_titles[i]}", key=f"copy_{i}", on_click=st.write, args=(section_content.strip(),))


    # --- Download Button ---
    st.markdown("---")
    st.subheader("Export Your Content")
    
    # Create the Word document in memory
    docx_buffer = create_word_document(st.session_state.generated_content)
    
    st.download_button(
        label="📥 Download as Word Document (.docx)",
        data=docx_buffer,
        file_name=f"{product_name.replace(' ', '_')}_marketing_content.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )